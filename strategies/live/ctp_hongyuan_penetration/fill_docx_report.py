#!/usr/bin/env python
"""Fill DOCX report template with Hongyuan penetration test results.

This module generates a Word document report for the Hongyuan CTP penetration
testing certification by filling a template with test results, screenshots,
and analysis.

Example:
    python fill_docx_report.py
"""
from __future__ import annotations

import json
import os
import re
import shutil
import textwrap
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

WORKSPACE_DIR = Path(__file__).resolve().parent
ROOT = WORKSPACE_DIR.parents[2]
TEMPLATE = WORKSPACE_DIR / "期货程序化交易系统功能测试过程记录报告模板.docx"
OUTPUT = WORKSPACE_DIR / "期货程序化交易系统功能测试过程记录报告.docx"
RESULTS_ROOT = WORKSPACE_DIR / "reports/latest"
SCREENSHOT_DIR = RESULTS_ROOT / "docx_log_screenshots"
VERSION_FILE = ROOT / "src/backend/app/__init__.py"
PACKAGE_DIR = ROOT / "src/backend"


def certification_order_symbol() -> str:
    """Return the contract used by the current Hongyuan certification run."""
    return (
        os.getenv("HONGYUAN_ORDER_SYMBOL")
        or os.getenv("HONGYUAN_TICK_SYMBOL")
        or "rb2610"
    ).strip()


SECTION_MAP = {
    2: ["C01"],
    3: ["T01", "T02", "T03"],
    4: ["M01", "M02", "M03"],
    5: ["M04", "M05"],
    6: ["O01", "O02", "O03"],
    7: ["TH01", "TH02", "TH03", "TH04", "TH05", "TH06"],
    8: ["V01", "V02", "V03"],
    9: ["E01", "E02", "E03"],
    10: ["EM01", "EM02", "EM03"],
    11: ["B01", "B02"],
    12: ["L01", "L02", "L03", "L04"],
}


JUDGEMENT_TEXT = {
    2: "通过。系统已完成柜台认证并成功登录测试账号。",
    3: "通过。系统已验证开仓、平仓、撤单基础交易功能正常。",
    4: "通过。系统已验证连接成功、断开显示及重连成功监测功能正常。",
    5: "通过。系统已验证报单笔数与撤单笔数统计功能正常。",
    6: "通过。系统已验证重复报单监测功能正常。",
    7: "通过。系统已验证阈值设置与达到或超过阈值后的预警功能正常。",
    8: "通过。系统已验证错误交易指令检查与拒绝功能正常。",
    9: "通过。系统已验证错误提示信息接收与展示功能正常。",
    10: "通过。系统已验证限制交易、暂停策略、强制退出等暂停交易功能正常。",
    11: "通过。系统已验证多笔部分成交报单与多笔已报单的批量撤单功能正常。",
    12: "通过。系统已验证交易日志、系统运行日志、监测日志、错误日志记录功能正常。",
}


REMARKS_TEXT = {
    11: "N/A（B01 已通过本地可控 Partial 状态构造结合 broker.batch_cancel() 正式路径完成稳定化验证）",
}


CASE_SEQUENCE = [
    "C01",
    "T01", "T02", "T03",
    "M01", "M02", "M03", "M04", "M05",
    "O01", "O02", "O03",
    "TH01", "TH02", "TH03", "TH04", "TH05", "TH06",
    "V01", "V02", "V03",
    "E01", "E02", "E03",
    "EM01", "EM02", "EM03",
    "B01", "B02",
    "L01", "L02", "L03", "L04",
]


CATEGORY_BY_CASE = {
    "C01": "连通性",
    "T01": "基础交易功能",
    "T02": "基础交易功能",
    "T03": "基础交易功能",
    "M01": "系统连接异常监测功能",
    "M02": "系统连接异常监测功能",
    "M03": "系统连接异常监测功能",
    "M04": "报撤单笔数监测功能",
    "M05": "报撤单笔数监测功能",
    "O01": "重复报单监测功能",
    "O02": "重复报单监测功能",
    "O03": "重复报单监测功能",
    "TH01": "阈值设置及预警功能",
    "TH02": "阈值设置及预警功能",
    "TH03": "阈值设置及预警功能",
    "TH04": "阈值设置及预警功能",
    "TH05": "阈值设置及预警功能",
    "TH06": "阈值设置及预警功能",
    "V01": "交易指令检查功能",
    "V02": "交易指令检查功能",
    "V03": "交易指令检查功能",
    "E01": "错误提示功能",
    "E02": "错误提示功能",
    "E03": "错误提示功能",
    "EM01": "暂停交易功能",
    "EM02": "暂停交易功能",
    "EM03": "暂停交易功能",
    "B01": "批量撤单功能",
    "B02": "批量撤单功能",
    "L01": "日志记录功能",
    "L02": "日志记录功能",
    "L03": "日志记录功能",
    "L04": "日志记录功能",
}


PREFERRED_LOGS_BY_CASE = {
    "C01": ["stdout.log", "logs/system.log"],
    "T01": ["stdout.log", "logs/order.log", "logs/system.log"],
    "T02": ["stdout.log", "logs/order.log", "logs/system.log"],
    "T03": ["stdout.log", "logs/order.log", "logs/system.log"],
    "M01": ["stdout.log", "logs/system.log"],
    "M02": ["stdout.log", "logs/system.log"],
    "M03": ["stdout.log", "logs/system.log"],
    "M04": ["stdout.log", "logs/monitor.log"],
    "M05": ["stdout.log", "logs/monitor.log"],
    "O01": ["stdout.log", "logs/monitor.log"],
    "O02": ["stdout.log", "logs/monitor.log"],
    "O03": ["stdout.log", "logs/monitor.log"],
    "TH01": ["stdout.log", "logs/monitor.log"],
    "TH02": ["stdout.log", "logs/monitor.log"],
    "TH03": ["stdout.log", "logs/monitor.log"],
    "TH04": ["stdout.log", "logs/monitor.log"],
    "TH05": ["stdout.log", "logs/monitor.log"],
    "TH06": ["stdout.log", "logs/monitor.log"],
    "V01": ["stdout.log", "logs/error.log"],
    "V02": ["stdout.log", "logs/error.log"],
    "V03": ["stdout.log", "logs/error.log"],
    "E01": ["stdout.log", "logs/error.log"],
    "E02": ["stdout.log", "logs/error.log"],
    "E03": ["stdout.log", "logs/error.log"],
    "EM01": ["stdout.log", "logs/system.log"],
    "EM02": ["stdout.log", "logs/system.log"],
    "EM03": ["stdout.log", "logs/system.log"],
    "B01": ["stdout.log", "logs/monitor.log", "logs/order.log"],
    "B02": ["stdout.log", "logs/monitor.log", "logs/order.log"],
    "L01": ["stdout.log", "logs/signal.log", "logs/order.log"],
    "L02": ["stdout.log", "logs/system.log"],
    "L03": ["stdout.log", "logs/monitor.log"],
    "L04": ["stdout.log", "logs/error.log"],
}


def load_results() -> dict[str, dict]:
    """Load all test results from the results directory.

    Returns:
        Dictionary mapping case IDs to their result data.

    Raises:
        RuntimeError: If not exactly 33 result files are found.
    """
    results = {}
    for child in RESULTS_ROOT.iterdir():
        result_path = child / "result.json"
        if child.is_dir() and result_path.exists():
            data = json.loads(result_path.read_text(encoding="utf-8"))
            results[data["case_id"]] = data
    if len(results) != 33:
        raise RuntimeError(f"Expected 33 result.json files, got {len(results)}")
    return results


def load_version() -> str:
    """Extract version string from version.py.

    Returns:
        Version string or "待补充" if not found.
    """
    if not VERSION_FILE.is_file():
        return "待补充"
    version_text = VERSION_FILE.read_text(encoding="utf-8")
    match = re.search(r'__version__\s*=\s*"([^"]+)"', version_text)
    return match.group(1) if match else "待补充"


def load_account_info() -> tuple[str, str]:
    """Load the local report identity without reading a credentials file.

    Returns:
        Tuple of (account_id, password).

    Raises:
        RuntimeError: If the local Hongyuan account ID is not configured.
    """
    account_id = os.getenv("HONGYUAN_USER_ID") or os.getenv("hongyuan_user_id")
    if not account_id:
        raise RuntimeError("请设置 HONGYUAN_USER_ID 后再生成 DOCX 报告")
    app_id = os.getenv("HONGYUAN_APP_ID", "client_wtyj_1.0.9.9")
    return account_id, app_id


def fmt(dt: datetime) -> str:
    """Format datetime as YYYY-MM-DD HH:MM:SS string."""
    return dt.strftime("%Y-%m-%d %H:%M:%S")


def fmt_range(results: dict[str, dict], ids: list[str]) -> str:
    """Format a time range for a set of test case IDs.

    Args:
        results: Dictionary of test results.
        ids: List of case IDs to format.

    Returns:
        Formatted time range string.
    """
    rows = [results[cid] for cid in ids]
    start = min(datetime.fromisoformat(row["started_at"]) for row in rows)
    end = max(datetime.fromisoformat(row["finished_at"]) for row in rows)
    return f"{fmt(start)} - {fmt(end)}"


def summary_line(results: dict[str, dict], ids: list[str]) -> str:
    """Generate a summary line of case statuses.

    Args:
        results: Dictionary of test results.
        ids: List of case IDs to include.

    Returns:
        Summary string with case IDs and statuses.
    """
    return "；".join(f"{cid}={results[cid]['status']}" for cid in ids)


def count_statuses(rows: list[dict]) -> dict[str, int]:
    """Count occurrences of each status in test results.

    Args:
        rows: List of result dictionaries.

    Returns:
        Dictionary with PASS, FAIL, BLOCKED counts.
    """
    counts = {"PASS": 0, "FAIL": 0, "BLOCKED": 0}
    for row in rows:
        counts[row["status"]] = counts.get(row["status"], 0) + 1
    return counts


def overall_conclusion(results: dict[str, dict]) -> str:
    """Generate overall conclusion text based on all test results.

    Args:
        results: Dictionary of test results.

    Returns:
        Conclusion text summarizing overall test outcome.
    """
    rows = [results[cid] for cid in CASE_SEQUENCE]
    counts = count_statuses(rows)
    total = len(rows)
    symbol = certification_order_symbol()
    if counts["PASS"] == total and counts["FAIL"] == 0 and counts["BLOCKED"] == 0:
        return f"宏源期货穿透式认证，测试合约 {symbol}，{total}/{total} PASS。"

    blocked_ids = [row["case_id"] for row in rows if row["status"] == "BLOCKED"]
    fail_ids = [row["case_id"] for row in rows if row["status"] == "FAIL"]
    parts = [
        f"宏源期货穿透式认证，测试合约 {symbol}，共 {total} 项",
        f"PASS {counts['PASS']} 项",
        f"BLOCKED {counts['BLOCKED']} 项",
        f"FAIL {counts['FAIL']} 项",
    ]
    text = "，".join(parts) + "。"
    if blocked_ids:
        text += f" 阻塞项：{'、'.join(blocked_ids)}。"
    if fail_ids:
        text += f" 失败项：{'、'.join(fail_ids)}。"
    return text


def build_section_judgement(table_index: int, results: dict[str, dict], ids: list[str]) -> str:
    """Build judgement text for a test section based on results.

    Args:
        table_index: Table index for looking up judgement text.
        results: Dictionary of test results.
        ids: List of case IDs in this section.

    Returns:
        Judgement text based on pass/fail status.
    """
    rows = [results[cid] for cid in ids]
    counts = count_statuses(rows)
    if counts["FAIL"] == 0 and counts["BLOCKED"] == 0:
        return JUDGEMENT_TEXT[table_index]
    if counts["PASS"] > 0:
        return (
            f"部分通过。本项共 {len(ids)} 个测试点，PASS {counts['PASS']} 个，"
            f"BLOCKED {counts['BLOCKED']} 个，FAIL {counts['FAIL']} 个。"
        )
    return (
        f"未通过。本项共 {len(ids)} 个测试点，PASS {counts['PASS']} 个，"
        f"BLOCKED {counts['BLOCKED']} 个，FAIL {counts['FAIL']} 个。"
    )


def build_section_remark(table_index: int, results: dict[str, dict], ids: list[str]) -> str:
    """Build remark text for a test section.

    Args:
        table_index: Table index for looking up base remark.
        results: Dictionary of test results.
        ids: List of case IDs in this section.

    Returns:
        Remark text including any failure reasons.
    """
    parts = []
    base_remark = REMARKS_TEXT.get(table_index, "").strip()
    if base_remark:
        parts.append(base_remark)

    nonpass_rows = [results[cid] for cid in ids if results[cid]["status"] != "PASS"]
    if nonpass_rows:
        parts.append(
            "未通过/阻塞项："
            + "；".join(
                f"{row['case_id']}={row['status']}（{row.get('failure_reason') or row.get('next_action') or '详见日志'}）"
                for row in nonpass_rows
            )
        )

    return "\n".join(parts)


def package_size_mb(path: Path) -> str:
    """Calculate total size of a directory in megabytes.

    Args:
        path: Directory path to calculate size for.

    Returns:
        Size string formatted as MB with 2 decimal places.
    """
    total = 0
    for child in path.rglob("*"):
        if child.is_file():
            total += child.stat().st_size
    return f"{total / 1024 / 1024:.2f}"


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """Load a font for image rendering.

    Args:
        size: Font size to use.

    Returns:
        FreeType font or default font if none found.
    """
    font_candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for candidate in font_candidates:
        if Path(candidate).exists():
            try:
                return ImageFont.truetype(candidate, size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def extract_non_empty_lines(path: Path) -> list[str]:
    """Extract non-empty lines from a text file.

    Args:
        path: Path to the text file.

    Returns:
        List of non-empty stripped lines.
    """
    if not path.exists():
        return []
    text = path.read_text(encoding="utf-8", errors="ignore").replace("\r\n", "\n")
    return [line.rstrip() for line in text.splitlines() if line.strip()]


def select_log_source(case_dir: Path, case_id: str) -> tuple[str, list[str]]:
    """Select the best log source for a test case.

    Args:
        case_dir: Directory containing case results.
        case_id: Case identifier.

    Returns:
        Tuple of (log_path, log_lines).
    """
    preferences = PREFERRED_LOGS_BY_CASE.get(case_id, ["stdout.log", "logs/system.log"])
    for rel_path in preferences:
        source_path = case_dir / rel_path
        lines = extract_non_empty_lines(source_path)
        if not lines:
            continue
        if rel_path.endswith("stdout.log"):
            return rel_path, lines[-12:]
        return rel_path, lines[:10]

    for source_path in sorted(case_dir.rglob("*.log")):
        lines = extract_non_empty_lines(source_path)
        if lines:
            rel_path = str(source_path.relative_to(case_dir))
            return rel_path, lines[:10]

    return "无可用日志", ["无可用日志内容"]


def wrap_line(line: str, width: int = 58) -> list[str]:
    """Wrap a line of text to a specified width.

    Args:
        line: Text line to wrap.
        width: Maximum line width.

    Returns:
        List of wrapped lines.
    """
    return textwrap.wrap(line, width=width, break_long_words=True, break_on_hyphens=False) or [""]


def render_log_screenshot(case_id: str, lines: list[str]) -> Path:
    """Render log lines as a PNG screenshot image.

    Args:
        case_id: Case identifier for the screenshot filename.
        lines: Log lines to render.

    Returns:
        Path to the generated screenshot file.
    """
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    body_font = load_font(22)
    wrapped_lines = []
    for line in lines:
        wrapped_lines.extend(wrap_line(line))

    width = 1400
    margin = 40
    line_height = 34
    top_bar_height = 46
    footer_height = 28
    height = top_bar_height + footer_height + margin * 2 + line_height * max(len(wrapped_lines), 1)
    image = Image.new("RGB", (width, height), color=(22, 27, 34))
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((14, 14, width - 14, height - 14), radius=18, fill=(11, 15, 20), outline=(74, 84, 98), width=2)
    draw.rounded_rectangle((14, 14, width - 14, 14 + top_bar_height), radius=18, fill=(26, 31, 39))
    draw.ellipse((34, 28, 48, 42), fill=(255, 95, 86))
    draw.ellipse((58, 28, 72, 42), fill=(255, 189, 46))
    draw.ellipse((82, 28, 96, 42), fill=(39, 201, 63))

    y = top_bar_height + margin
    for line in wrapped_lines:
        draw.text((margin, y), line, font=body_font, fill=(215, 223, 232))
        y += line_height

    output = SCREENSHOT_DIR / f"{case_id}.png"
    image.save(output)
    return output


def insert_paragraph_after(anchor, text: str = "", style: str | None = None, align=None) -> Paragraph:
    """Insert a new paragraph after an anchor element.

    Args:
        anchor: The anchor element to insert after.
        text: Optional text content for the paragraph.
        style: Optional style name for the paragraph.
        align: Optional alignment for the paragraph.

    Returns:
        The newly created paragraph.
    """
    new_p = OxmlElement("w:p")
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_inline_log_screenshots(doc: Document, results: dict[str, dict]) -> None:
    """Add inline log screenshots to the document for each test case.

    Args:
        doc: The Document to add screenshots to.
        results: Dictionary of test results.
    """
    if SCREENSHOT_DIR.exists():
        shutil.rmtree(SCREENSHOT_DIR)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)

    for table_index, case_ids in SECTION_MAP.items():
        anchor = doc.tables[table_index]
        title_para = insert_paragraph_after(anchor, "测试点执行日志截图", style="Heading 3")
        cursor = title_para

        for case_id in case_ids:
            case_result = results[case_id]
            case_dir = RESULTS_ROOT / case_id
            rel_log_path, raw_lines = select_log_source(case_dir, case_id)
            image_path = render_log_screenshot(case_id, raw_lines)
            info_lines = [
                f"日志文件：reports/latest/{case_id}/{rel_log_path}",
                f"验证结果：{case_result['status']}",
                f"验证时间：{case_result['started_at']} ~ {case_result['finished_at']}",
            ]
            failure_reason = case_result.get("failure_reason") or case_result.get("next_action")
            if failure_reason:
                info_lines.append(f"原因：{failure_reason}")

            case_para = insert_paragraph_after(cursor, f"{case_id} {case_result['case_name']}", style="Heading 4")
            info_para = insert_paragraph_after(
                case_para,
                "\n".join(info_lines),
            )
            picture_para = insert_paragraph_after(info_para, align=WD_ALIGN_PARAGRAPH.CENTER)
            picture_para.add_run().add_picture(str(image_path), width=Inches(6.5))
            spacer = insert_paragraph_after(picture_para, "")
            cursor = spacer


def fill_doc() -> Path:
    """Fill the DOCX template with test results and generate the report.

    Returns:
        Path to the generated report file.
    """
    if not TEMPLATE.is_file():
        raise FileNotFoundError(
            f"缺少报告模板: {TEMPLATE}。请将宏源期货认可的 DOCX 模板放入工作区后重试。"
        )

    results = load_results()
    version = load_version()
    account_id, app_id = load_account_info()

    all_started = [datetime.fromisoformat(row["started_at"]) for row in results.values() if row.get("started_at")]
    all_finished = [datetime.fromisoformat(row["finished_at"]) for row in results.values() if row.get("finished_at")]
    start_dt = min(all_started)
    finish_dt = max(all_finished)

    doc = Document(str(TEMPLATE))

    doc.tables[0].cell(0, 1).text = "backtrader 项目组"
    doc.tables[0].cell(1, 1).text = "待补充"
    doc.tables[0].cell(2, 1).text = account_id
    doc.tables[0].cell(3, 1).text = f"{fmt(start_dt)} - {fmt(finish_dt)}"
    doc.tables[0].cell(4, 1).text = overall_conclusion(results)

    doc.tables[1].cell(0, 1).text = "backtrader 期货程序化交易系统"
    doc.tables[1].cell(1, 1).text = version
    doc.tables[1].cell(2, 1).text = f"{package_size_mb(PACKAGE_DIR)} MB"
    doc.tables[1].cell(3, 1).text = "backtrader"
    doc.tables[1].cell(4, 1).text = "待补充"
    doc.tables[1].cell(5, 1).text = "直连投资者"
    doc.tables[1].cell(6, 1).text = app_id
    doc.tables[1].cell(7, 1).text = "CTP期货仿真（上海唐银，BrokerID=3070）"

    for table_index, ids in SECTION_MAP.items():
        table = doc.tables[table_index]
        process_text = table.cell(2, 1).text.strip()
        process_text += f"\n执行时间：{fmt_range(results, ids)}\n执行结果：{summary_line(results, ids)}"
        table.cell(2, 1).text = process_text
        table.cell(3, 1).text = build_section_judgement(table_index, results, ids)
        remark_text = build_section_remark(table_index, results, ids)
        if remark_text:
            table.cell(4, 1).text = remark_text

    add_inline_log_screenshots(doc, results)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    path = fill_doc()
    print(path)
